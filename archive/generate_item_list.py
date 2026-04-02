import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 데이터 로드 ───
with open('data/quiz_items.json', encoding='utf-8') as f:
    items = json.load(f)

# ─── 퀴즈 포함 여부 (index.html 기준) ───
ITEM_HIDE_CATS = {'all-machines', 'scarves', 'plates', 'jewels', 'memories', 'z-crystals'}

def in_quiz(item):
    return item['category'] not in ITEM_HIDE_CATS

# ─── 엔트리 도구 포함 여부 (entry.html 기준) ───
ALLOWED_CATS = {
    'held-items', 'choice', 'scarves', 'bad-held-items', 'in-a-pinch',
    'species-specific', 'type-enhancement', 'type-protection',
    'mega-stones', 'memories', 'plates',
    'spelunking', 'jewels',
}
HIDDEN_ITEMS = {
    '다홍꿀', '진노랑꿀', '연분홍꿀', '보라꿀', '나무열매쥬스',
    '검정비드로', '하양비드로', '에나비꼬리', '삐삐인형',
    '실버스프레이', '골드스프레이', '벌레회피스프레이', '동굴탈출로프',
    '심해의이빨', '심해의비늘',
}

def in_entry(item):
    if item['ko'] in HIDDEN_ITEMS:
        return False
    return item['category'] in ALLOWED_CATS or '열매' in item['ko']

# ─── 정렬: 카테고리 → ID ───
items_sorted = sorted(items, key=lambda x: (x['category'], x['id']))

# ─── docx 생성 ───
doc = Document()

# 기본 폰트 설정
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(9)

# 제목
title = doc.add_heading('Pokénova 도구 전체 목록', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'총 {len(items_sorted)}개 항목 | 퀴즈O: {sum(in_quiz(i) for i in items_sorted)}개 | 엔트리O: {sum(in_entry(i) for i in items_sorted)}개')

doc.add_paragraph('')

# ─── 테이블 생성 ───
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# 컬럼 너비 설정
col_widths = [Cm(3.5), Cm(2.5), Cm(8.0), Cm(2.0), Cm(2.5)]
for i, width in enumerate(col_widths):
    for cell in table.columns[i].cells:
        cell.width = width

# 헤더
hdr = table.rows[0].cells
headers = ['카테고리', '이름 (한국어)', '설명', '퀴즈여부', '엔트리도구여부']
for i, h in enumerate(headers):
    hdr[i].text = h
    run = hdr[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9)
    # 헤더 배경색
    tc = hdr[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2C2C2C')
    tcPr.append(shd)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# 카테고리별 색상 매핑
CAT_COLORS = {
    'held-items':       'F3E5F5',
    'choice':           'EDE7F6',
    'bad-held-items':   'FFEBEE',
    'in-a-pinch':       'FFF3E0',
    'species-specific': 'E8F5E9',
    'type-enhancement': 'E3F2FD',
    'type-protection':  'E1F5FE',
    'mega-stones':      'FCE4EC',
    'memories':         'F1F8E9',
    'plates':           'FFF8E1',
    'spelunking':       'EFEBE9',
    'jewels':           'F9FBE7',
    'scarves':          'E0F2F1',
}

prev_cat = None

for item in items_sorted:
    row = table.add_row()
    cells = row.cells

    cat = item['category']
    ko  = item.get('ko', '')
    desc = item.get('desc', '')
    quiz_yn   = 'O' if in_quiz(item)  else 'X'
    entry_yn  = 'O' if in_entry(item) else 'X'

    cells[0].text = cat
    cells[1].text = ko
    cells[2].text = desc
    cells[3].text = quiz_yn
    cells[4].text = entry_yn

    # 행 배경색
    fill_color = CAT_COLORS.get(cat, 'FFFFFF')
    for cell in cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(8.5)


# 저장
out = 'item_list.docx'
doc.save(out)
print(f'저장 완료: {out}')
print(f'총 {len(items_sorted)}개 | 퀴즈O: {sum(in_quiz(i) for i in items_sorted)} | 엔트리O: {sum(in_entry(i) for i in items_sorted)}')
