import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('/home/soondoree07/pokemon_project_1/data/quiz_items.json', encoding='utf-8') as f:
    items = json.load(f)

# 숨김 카테고리 제외
HIDE_CATS = {'all-machines', 'scarves', 'plates', 'jewels', 'memories'}

# 표시 모드
IMAGE_ONLY = {'mega-stones', 'species-specific'}
DESC_ONLY  = {'bad-held-items', 'type-enhancement', 'choice', 'held-items', 'evolution'}

# 카테고리 한국어 이름
CAT_NAMES = {
    'standard-balls':   '몬스터볼 (기본)',
    'special-balls':    '몬스터볼 (특수)',
    'apricorn-balls':   '몬스터볼 (나무열매)',
    'healing':          '회복약',
    'status-cures':     '상태이상 치료',
    'revival':          '부활',
    'pp-recovery':      'PP 회복',
    'vitamins':         '비타민',
    'medicine':         '치료 열매',
    'picky-healing':    '쓴맛 열매',
    'effort-drop':      '노력치 감소 열매',
    'other':            '기타 열매',
    'type-protection':  '타입반감 열매',
    'in-a-pinch':       '피지 열매',
    'effort-training':  '파워계열 (노력치)',
    'training':         '훈련 도구',
    'spelunking':       '던전 도구',
    'flutes':           '비드로',
    'stat-boosts':      '배틀 아이템',
    'mega-stones':      '메가스톤 ★이미지만',
    'species-specific': '종족전용 ★이미지만',
    'bad-held-items':   '방해 도구 ★설명만',
    'type-enhancement': '타입강화 도구 ★설명만',
    'choice':           '구애 3종 ★설명만',
    'held-items':       '지닌 도구 ★설명만',
    'evolution':        '진화 아이템 ★설명만',
}

# 카테고리 정렬 순서
CAT_ORDER = [
    'mega-stones', 'species-specific',
    'evolution', 'held-items', 'bad-held-items', 'type-enhancement', 'choice',
    'standard-balls', 'special-balls', 'apricorn-balls',
    'healing', 'status-cures', 'revival', 'pp-recovery',
    'vitamins',
    'medicine', 'picky-healing', 'effort-drop', 'other', 'type-protection', 'in-a-pinch',
    'effort-training', 'training', 'spelunking', 'flutes', 'stat-boosts',
]

quiz_items = [x for x in items if x['category'] not in HIDE_CATS]

# 카테고리별 그룹
from collections import defaultdict
grouped = defaultdict(list)
for x in quiz_items:
    grouped[x['category']].append(x)

doc = Document()
title = doc.add_heading('도구 퀴즈 아이템 목록', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'총 {len(quiz_items)}개 아이템 (숨김 제외)\n')

for cat in CAT_ORDER:
    if cat not in grouped:
        continue
    cat_items = grouped[cat]
    mode = '이미지만' if cat in IMAGE_ONLY else ('설명만→이미지공개' if cat in DESC_ONLY else '이미지+설명')
    cat_label = CAT_NAMES.get(cat, cat)

    # 카테고리 헤더
    h = doc.add_heading(f'{cat_label}  [{mode}]  ({len(cat_items)}개)', level=2)

    for x in cat_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)

        run_name = p.add_run(f'[{x["id"]}] {x["ko"]}')
        run_name.bold = True
        run_name.font.size = Pt(11)

        run_en = p.add_run(f'  ({x["en"]})')
        run_en.font.size = Pt(9)
        run_en.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.4)
        r2 = p2.add_run(x.get('desc', '설명 없음'))
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

    doc.add_paragraph()

out = '/home/soondoree07/pokemon_project_1/도구퀴즈_아이템목록.docx'
doc.save(out)
print(f'저장 완료: {out}')
print(f'총 {len(quiz_items)}개')
for cat in CAT_ORDER:
    if cat in grouped:
        print(f'  {CAT_NAMES.get(cat,cat)}: {len(grouped[cat])}개')
