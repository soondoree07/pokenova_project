"""
한국어 미번역 기술 목록 → docx 파일 생성
data/missing_ko_moves.json → 미번역_기술_목록.docx
"""
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('data/missing_ko_moves.json', encoding='utf-8') as f:
    moves = json.load(f)

needs_name = [m for m in moves if m.get('needs_ko')]
needs_desc = [m for m in moves if not m.get('needs_ko')]

doc = Document()

# 제목
title = doc.add_heading('미번역 기술 목록 (SV 9세대)', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'총 {len(moves)}개 기술 번역 필요 — 이름 없음: {len(needs_name)}개 / 설명 없음: {len(needs_desc)}개')
doc.add_paragraph('')

# ── 섹션 1: 이름 없는 기술 ──
doc.add_heading(f'1. 한국어 이름 없음 ({len(needs_name)}개)', level=2)
doc.add_paragraph('영문 이름 기준으로 수록되어 있음. 한국어 이름 + 설명 번역 필요.')
doc.add_paragraph('')

table1 = doc.add_table(rows=1, cols=6)
table1.style = 'Table Grid'
hdr = table1.rows[0].cells
hdr[0].text = 'ID'
hdr[1].text = '영문명 (en)'
hdr[2].text = '타입'
hdr[3].text = '분류'
hdr[4].text = 'PP / 위력 / 명중'
hdr[5].text = '영문 설명'
for cell in hdr:
    for p in cell.paragraphs:
        run = p.runs[0] if p.runs else p.add_run(cell.text)
        run.bold = True

for m in sorted(needs_name, key=lambda x: x['id']):
    row = table1.add_row().cells
    row[0].text = str(m['id'])
    row[1].text = m['en']
    row[2].text = m.get('type', '')
    row[3].text = m.get('class', '')
    pp  = m.get('pp') or '-'
    pow_ = m.get('power') or '-'
    acc = m.get('accuracy') or '-'
    row[4].text = f"{pp} / {pow_} / {acc}"
    row[5].text = m.get('desc', '')

doc.add_paragraph('')

# ── 섹션 2: 설명 없는 기술 ──
doc.add_heading(f'2. 한국어 설명 없음 ({len(needs_desc)}개)', level=2)
doc.add_paragraph('이름은 있으나 한국어 설명 텍스트가 누락됨. 설명 번역만 필요.')
doc.add_paragraph('')

table2 = doc.add_table(rows=1, cols=6)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'ID'
hdr2[1].text = '한국어명 (ko)'
hdr2[2].text = '영문명 (en)'
hdr2[3].text = '타입'
hdr2[4].text = '분류'
hdr2[5].text = '영문 설명'
for cell in hdr2:
    for p in cell.paragraphs:
        run = p.runs[0] if p.runs else p.add_run(cell.text)
        run.bold = True

for m in sorted(needs_desc, key=lambda x: x['id']):
    row = table2.add_row().cells
    row[0].text = str(m['id'])
    row[1].text = m.get('ko', '')
    row[2].text = m.get('en', '')
    row[3].text = m.get('type', '')
    row[4].text = m.get('class', '')
    row[5].text = m.get('desc', '')

out = '미번역_기술_목록.docx'
doc.save(out)
print(f'✅ {out} 저장 완료 ({len(moves)}개 기술)')
