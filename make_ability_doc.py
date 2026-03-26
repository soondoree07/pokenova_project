import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('/home/soondoree07/pokemon_project_1/data/quiz_abilities.json', encoding='utf-8') as f:
    abilities = json.load(f)

with open('/home/soondoree07/pokemon_project_1/data/quiz_pokemon.json', encoding='utf-8') as f:
    pokemon_data = json.load(f)

# 특성 이름 기준으로 일반/숨겨진 보유 포켓몬 맵 구축
# ability_ko → { normal: [], hidden: [] }
ab_map = {}

def add_pokemon(ability_ko, pokemon_ko, is_hidden):
    if ability_ko not in ab_map:
        ab_map[ability_ko] = {'normal': [], 'hidden': []}
    key = 'hidden' if is_hidden else 'normal'
    if pokemon_ko not in ab_map[ability_ko][key]:
        ab_map[ability_ko][key].append(pokemon_ko)

for p in pokemon_data:
    for ab in p.get('abilities', []):
        if isinstance(ab, dict) and ab.get('ko'):
            add_pokemon(ab['ko'], p['ko'], ab.get('hidden', False))
    for form in p.get('forms', []):
        form_name = form.get('name_ko', p['ko'])
        for ab in form.get('abilities', []):
            if isinstance(ab, dict) and ab.get('ko'):
                add_pokemon(ab['ko'], form_name, ab.get('hidden', False))

doc = Document()

title = doc.add_heading('포켓몬 특성 목록', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'총 {len(abilities)}개 특성\n')

abilities = sorted(abilities, key=lambda a: a.get('ko', ''))

for a in abilities:
    aid = a.get('id', '')
    ko = a.get('ko', '')
    en = a.get('en', '')
    desc = a.get('desc', '설명 없음')

    entry = ab_map.get(ko, {'normal': [], 'hidden': []})
    normal_list = entry['normal']
    hidden_list = entry['hidden']

    # 전체 포켓몬 목록 (숨겨진 특성은 * 표시)
    combined = normal_list + [f'{pk}*' for pk in hidden_list]

    # 특성 이름 (굵게)
    p = doc.add_paragraph()
    run = p.add_run(f'[{aid}] {ko}')
    run.bold = True
    run.font.size = Pt(12)
    run_en = p.add_run(f'  ({en})')
    run_en.font.size = Pt(10)
    run_en.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 설명
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.add_run(f'설명: {desc}').font.size = Pt(10)

    # 보유 포켓몬 (숨겨진 특성은 * 표시)
    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Inches(0.2)
    r3 = p3.add_run(f'보유 포켓몬: {", ".join(combined) if combined else "없음"}')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)

    if hidden_list:
        p4 = doc.add_paragraph()
        p4.paragraph_format.left_indent = Inches(0.2)
        r4 = p4.add_run('* 표시는 숨겨진 특성')
        r4.font.size = Pt(9)
        r4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        r4.italic = True

    doc.add_paragraph()

out_path = '/home/soondoree07/pokemon_project_1/포켓몬_특성_목록.docx'
doc.save(out_path)
print(f'저장 완료: {out_path}')
print(f'총 {len(abilities)}개 특성')
